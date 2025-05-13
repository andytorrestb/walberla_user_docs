from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

# Lit-review data with 25 entries
entries = [
    {
        "title": "TechRep_2004-02.pdf",
        "summary": "Evaluates the parallel performance of large-scale LBM applications on various supercomputing architectures. Emphasizes the importance of efficient CPUs and high-speed communication for optimal LBM performance.",
        "author": "Thomas Pohl et al.",
        "date": "February 2005"
    },
    {
        "title": "TechRep_2005-07.pdf",
        "summary": "Explores adaptive parameterization for free surface flows using LBM, with emphasis on Smagorinsky turbulence modeling and adaptive timesteps.",
        "author": "Gerhard Wellein et al.",
        "date": "July 2005"
    },
    {
        "title": "wellein2006.pdf",
        "summary": "Studies how data layout and architecture-specific optimizations affect single-processor performance of LBM kernels. Finds that vector systems can significantly outperform standard CPUs.",
        "author": "Gerhard Wellein et al.",
        "date": "2006"
    },
    {
        "title": "Faisal_MT_2006.pdf",
        "summary": "A master’s thesis on data transfer between non-matching meshes in multiphysics simulations using radial basis functions and isoparametric mapping.",
        "author": "Faisal et al.",
        "date": "2006"
    },
    {
        "title": "TechRep_2008-01.pdf",
        "summary": "Provides foundational work in data structures or algorithm design for LBM or multibody dynamics.",
        "author": "Unknown",
        "date": "2008"
    },
    {
        "title": "TechRep_2009-08.pdf",
        "summary": "Technical report likely on performance tuning or simulation studies in LBM or rigid body simulations.",
        "author": "Unknown",
        "date": "2009"
    },
    {
        "title": "TechRep_2009-10.pdf",
        "summary": "Introduces compile-time assertions to enforce type and interface consistency in C++ physics engine code for simulations.",
        "author": "U. Rüde et al.",
        "date": "2009"
    },
    {
        "title": "TechRep_2009-11.pdf",
        "summary": "Covers internal software engineering strategies used in waLBerla, possibly introducing custom data containers.",
        "author": "U. Rüde et al.",
        "date": "2009"
    },
    {
        "title": "TechRep_2009-18.pdf",
        "summary": "Introduces the Sandwich Pattern—a hybrid compile-time and runtime polymorphism design pattern used in the physics engine.",
        "author": "U. Rüde et al.",
        "date": "2009"
    },
    {
        "title": "TechRep_2009-19.pdf",
        "summary": "Explains setup and parallelization strategy of large-scale rigid multibody simulations using the pe physics engine.",
        "author": "U. Rüde et al.",
        "date": "2009"
    },
    {
        "title": "TechRep_2010-03.pdf",
        "summary": "Likely discusses performance strategies or a module addition to waLBerla or the physics engine in simulation environments.",
        "author": "U. Rüde et al.",
        "date": "2010"
    },
    {
        "title": "TechRep_2010-07.pdf",
        "summary": "Presents a performance model and GPU acceleration of LBM within waLBerla, using CPU-GPU heterogeneous nodes.",
        "author": "Christian Feichtinger et al.",
        "date": "2010"
    },
    {
        "title": "Gotz2010.pdf",
        "summary": "Describes a coupled LBM and rigid-body physics simulation framework that scales to 8192 cores. Demonstrates application in particle-laden flows.",
        "author": "Jan Götz et al.",
        "date": "2010"
    },
    {
        "title": "feichtinger2011.pdf",
        "summary": "Presents the WaLBerla framework for high-performance multiphysics simulations, focusing on modularity, scalability, and hybrid CPU-GPU execution.",
        "author": "Christian Feichtinger et al.",
        "date": "2011"
    },
    {
        "title": "Rauh_BT_2013.pdf",
        "summary": "Bachelor thesis on accurate LBM treatment of moving boundaries using triangulated surfaces, improving over bounce-back methods.",
        "author": "Matthias Rauh",
        "date": "2013"
    },
    {
        "title": "Rettinger_BT_2013.pdf",
        "summary": "Bachelor thesis evaluating BGK, TRT, and MRT collision models in LBM, including parameter selection and benchmark testing.",
        "author": "Christoph Rettinger",
        "date": "2013"
    },
    {
        "title": "popa2014.pdf",
        "summary": "Introduces regularized solution methods for linear complementarity problems (LCPs) with application to multibody dynamics. Proposes iterative solvers and analyzes their convergence behavior.",
        "author": "Constantin Popa et al.",
        "date": "2014"
    },
    {
        "title": "Krieg_BT_2019.pdf",
        "summary": "Bachelor thesis analyzing grid refinement strategies in LBM and validating results against known solutions like Couette flow.",
        "author": "Benedikt Krieg",
        "date": "2019"
    },
    {
        "title": "Holzer_MT_2020.pdf",
        "summary": "Master’s thesis focused on LBM simulation performance with complex geometries, advanced collision models, and GPU integration.",
        "author": "Markus Holzer",
        "date": "2020"
    },
    {
        "title": "Schwarzmeier_PASC_2021.pdf",
        "summary": "Poster on LBM-based DNS simulations through porous media using waLBerla with validation against turbulence models.",
        "author": "Christoph Schwarzmeier",
        "date": "2021"
    },
    {
        "title": "Diss_2017-Fattahi.pdf",
        "summary": "PhD dissertation (summary unavailable in search result) related to LBM or coupled physics simulations in waLBerla.",
        "author": "Ehsan Fattahi",
        "date": "2017"
    },
    {
        "title": "Schornbaum_DA_2010.pdf",
        "summary": "Diploma thesis on impulse-based rigid body collision response, including shared memory parallelization and stability analysis.",
        "author": "Florian Schornbaum",
        "date": "2010"
    },
    {
        "title": "Diss_2018-Schornbaum.pdf",
        "summary": "PhD thesis on adaptive mesh refinement (AMR) for extreme-scale LBM simulations using distributed block-based partitioning.",
        "author": "Florian Schornbaum",
        "date": "2018"
    },
    {
        "title": "Schornbaum_Paris_SIAM-PP_2016-04-15.pdf",
        "summary": "Presentation on adaptive LBM mesh refinement and block-based domain partitioning for extreme-scale supercomputing.",
        "author": "Florian Schornbaum",
        "date": "2016"
    },
    {
        "title": "vsc-18-cpr-v1.pdf",
        "summary": "Presentation of waLBerla applications in 3D LBM simulations for additive manufacturing and validation against industrial test cases.",
        "author": "Ulrich Rüde",
        "date": "2018"
    }
]


# Extract the year from the date string for sorting
def extract_year(date_str):
    match = re.search(r'\d{4}', date_str)
    return int(match.group()) if match else 0

# Sort entries by year ascending
sorted_entries = sorted(entries, key=lambda x: extract_year(x["date"]))

# Create Word document
doc = Document()
doc.add_heading('Document Summaries (Ordered by Date: Oldest to Newest)', 0)

# Create a table for each entry
for i, entry in enumerate(sorted_entries, start=1):
    doc.add_heading(f"{i}. {entry['title']}", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    table.style = 'Table Grid'

    table.cell(0, 0).text = "Author(s)"
    table.cell(0, 1).text = entry["author"]

    table.cell(1, 0).text = "Publication Date"
    table.cell(1, 1).text = entry["date"]

    table.cell(2, 0).text = "Summary"
    table.cell(2, 1).text = entry["summary"]

    doc.add_paragraph("")  # spacing

# Save the file
output_path = "Document_Summaries_Example.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
